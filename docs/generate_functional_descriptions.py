"""Generate Functional-Descriptions.docx for AdventureLog's critical functions."""
from docx import Document
from docx.shared import Pt
from pathlib import Path

OUT = Path(__file__).parent / "Functional-Descriptions.docx"

# Each tuple: (function name, input, processing, output)
DESCRIPTIONS = [
    (
        "Sign up a new user",
        "New username, new password, repeating password",
        "Checks if username is available (loops until a unique name is entered), "
        "checks if both passwords match (loops until they do), creates a new User "
        "object, registers it with State and Disk, and persists the change",
        "Success or error message; new user written to disk.yaml",
    ),
    (
        "Log in an existing user",
        "Username, password",
        "Iterates over the loaded users; if a name matches, the password is compared. "
        "Throws \"User does not exist\" if no name matches, or \"Incorrect Password\" "
        "if the name matches but the password does not. On success, sets the current "
        "user pointer in State",
        "Success (user is logged in and main menu is shown) or error message",
    ),
    (
        "Create a new log",
        "Log type (Hike or Cave), date, duration, area, "
        "type-specific fields (cave name and lead/SRT/rig flags for Cave; distance "
        "and weather for Hike), optional participant count and names, free-text note",
        "Validates the type and date are non-empty, builds either a CaveLog or "
        "HikeLog with the entered fields, attaches each named participant to the log, "
        "links the log to the current user, and triggers a save",
        "New log added to the user's log list and persisted to disk; "
        "confirmation or error message",
    ),
    (
        "View a log",
        "Index of the log selected from the log overview page",
        "Looks up the log on the current user's log list; if the index is in range, "
        "calls the log's polymorphic print() method (Cave or Hike specific). "
        "Then offers navigation back to the overview or main menu",
        "Full log details printed to the console, or an "
        "\"Log doesn't exist\" error message",
    ),
    (
        "Add participants to a log",
        "Number of participants, name of each participant",
        "Loops the requested number of times, prompting for each name and rejecting "
        "empty inputs. Each accepted name is appended to a list, then later turned "
        "into a Participant object attached to the parent log",
        "Participants added to the log and persisted with it on save",
    ),
    (
        "Save state to disk",
        "None (called automatically on clean exit, or explicitly after sign-up / "
        "log creation)",
        "Walks every user, then every log under each user, then every participant "
        "under each log, serialising each object into key/value lines via "
        "Disk::userToStr / logToStr / partToStr. Concatenates the buffers in "
        "parent-first order and rewrites disk.yaml with ios::trunc",
        "disk.yaml fully rewritten with the current in-memory state; "
        "any orphan logs or participants are silently dropped",
    ),
    (
        "Load save from disk",
        "Path to a save file (defaults to disk.yaml)",
        "Opens the file, parses each line into a key/value pair, splits the stream "
        "on \"---\" object dividers, then loads in child-first order: participants, "
        "then logs (wiring participants to their parent log by log-id), then users "
        "(wiring logs to their owner by owner-id). Restores the static ID counters "
        "to the largest ID seen per type",
        "Fully populated in-memory object graph; on failure an error is printed "
        "and the user is reprompted for a valid path",
    ),
]


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("Functional Description")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subrun = subtitle.add_run(
        "AdventureLog — critical functions covering authentication, "
        "log creation, viewing, and persistence."
    )
    subrun.italic = True

    doc.add_paragraph()

    for name, inputs, processing, output in DESCRIPTIONS:
        p = doc.add_paragraph()
        r = p.add_run("Function: ")
        r.bold = True
        p.add_run(name)

        p = doc.add_paragraph()
        r = p.add_run("Input: ")
        r.bold = True
        p.add_run(inputs)

        p = doc.add_paragraph()
        r = p.add_run("Processing: ")
        r.bold = True
        p.add_run(processing)

        p = doc.add_paragraph()
        r = p.add_run("Output: ")
        r.bold = True
        p.add_run(output)

        doc.add_paragraph()

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
